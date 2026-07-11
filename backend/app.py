from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_compress import compress
import PyPDF2
import pytesseract
from PIL import Image
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os
import re
from datetime import datetime
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from transformers import pipeline
import logging
from werkzeug.utils import secure_filename
import json

# Download required NLTK data
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

app = Flask(__name__)
CORS(app)
compress(app)

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize AI models
@limiter.limit("10 per minute")
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'OK',
        'timestamp': datetime.utcnow().isoformat(),
        'version': '1.0.0'
    })

# Text processing pipeline
text_processor = pipeline("text-classification", model="distilgpt2")

# Document type detection patterns
ACADEMIC_PATTERNS = {
    'introduction': ['introduction', 'intro', 'background', 'overview', 'preliminaries'],
    'methodology': ['methodology', 'methods', 'approach', 'procedure', 'experiment', 'design'],
    'results': ['results', 'findings', 'analysis', 'data', 'observations', 'evaluation'],
    'conclusion': ['conclusion', 'summary', 'discussion', 'implications', 'future work'],
    'references': ['references', 'bibliography', 'works cited', 'citations', 'literature review']
}

# Initialize stopwords for text processing
stop_words = set(stopwords.words('english'))

@limiter.limit("5 per minute")
@app.route('/api/process-text', methods=['POST'])
def process_text():
    """Process and structure text input"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data['text']
        doc_type = data.get('type', 'academic')
        
        # Process the text
        structured_content = process_text_content(text, doc_type)
        
        # Generate document
        document = generate_document(structured_content, doc_type)
        
        return jsonify({
            'success': True,
            'document': document,
            'structuredContent': structured_content
        })
        
    except Exception as e:
        logger.error(f"Error processing text: {str(e)}")
        return jsonify({'error': 'Failed to process text'}), 500

@limiter.limit("5 per minute")
@app.route('/api/process-file', methods=['POST'])
def process_file():
    """Process uploaded file (PDF, image, text)"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'File is required'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate file type
        allowed_extensions = {'pdf', 'png', 'jpg', 'jpeg', 'txt', 'docx'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        
        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Unsupported file type'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        temp_path = os.path.join('/tmp', filename)
        file.save(temp_path)
        
        # Extract text based on file type
        extracted_text = ''
        if file_ext == 'pdf':
            extracted_text = extract_text_from_pdf(temp_path)
        elif file_ext in ['png', 'jpg', 'jpeg']:
            extracted_text = extract_text_from_image(temp_path)
        elif file_ext == 'txt':
            extracted_text = extract_text_from_text_file(temp_path)
        elif file_ext == 'docx':
            extracted_text = extract_text_from_docx(temp_path)
        
        # Clean up temp file
        os.unlink(temp_path)
        
        # Process the extracted text
        doc_type = request.form.get('type', 'academic')
        structured_content = process_text_content(extracted_text, doc_type)
        
        # Generate document
        document = generate_document(structured_content, doc_type)
        
        return jsonify({
            'success': True,
            'originalFilename': file.filename,
            'document': document,
            'structuredContent': structured_content,
            'extractedText': extracted_text
        })
        
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        return jsonify({'error': 'Failed to process file'}), 500

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n\n'
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise

def extract_text_from_image(file_path):
    """Extract text from image using OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.error(f"Image OCR error: {str(e)}")
        raise

def extract_text_from_text_file(file_path):
    """Extract text from text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        logger.error(f"Text file read error: {str(e)}")
        raise

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise

def process_text_content(text, doc_type='academic'):
    """Process and structure text content using NLP"""
    try:
        # Tokenize text
        sentences = sent_tokenize(text)
        words = word_tokenize(text)
        
        # Remove stopwords
        filtered_words = [word for word in words if word.lower() not in stop_words]
        
        # Detect document structure
        structure = detect_document_structure(text)
        
        # Group related content
        grouped_content = group_related_content(sentences, structure)
        
        # Apply formatting based on type
        formatted_content = apply_formatting(grouped_content, doc_type)
        
        return {
            'originalText': text,
            'wordCount': len(filtered_words),
            'sentenceCount': len(sentences),
            'structure': structure,
            'groupedContent': grouped_content,
            'formattedContent': formatted_content
        }
        
    except Exception as e:
        logger.error(f"Text processing error: {str(e)}")
        raise

def detect_document_structure(text):
    """Detect document structure using NLP"""
    structure = {
        'sections': [],
        'headings': [],
        'subheadings': [],
        'detectedTypes': []
    }
    
    lower_text = text.lower()
    
    # Detect sections based on patterns
    for section_type, patterns in ACADEMIC_PATTERNS.items():
        for pattern in patterns:
            if pattern in lower_text:
                structure['detectedTypes'].append(section_type)
                structure['sections'].append({
                    'type': section_type,
                    'pattern': pattern,
                    'position': lower_text.index(pattern)
                })
    
    # Detect headings (capitalized lines)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        trimmed = line.strip()
        if (trimmed and len(trimmed) < 100 and 
            trimmed[0].isupper() and not trimmed.endswith('.') and
            not any(char in trimmed for char in '[](){}')):
            if len(trimmed) > 10:
                structure['headings'].append({
                    'text': trimmed,
                    'line': i,
                    'type': 'main'
                })
            else:
                structure['subheadings'].append({
                    'text': trimmed,
                    'line': i,
                    'type': 'sub'
                })
    
    return structure

def group_related_content(sentences, structure):
    """Group related content based on detected structure"""
    groups = {
        'introduction': [],
        'methodology': [],
        'results': [],
        'conclusion': [],
        'references': [],
        'general': []
    }
    
    for sentence in sentences:
        lower_sentence = sentence.lower()
        assigned = False
        
        # Assign to appropriate section based on keywords
        for section_type in structure['detectedTypes']:
            if section_type in lower_sentence and not assigned:
                groups[section_type].append(sentence)
                assigned = True
                break
        
        if not assigned:
            groups['general'].append(sentence)
    
    return groups

def apply_formatting(content, doc_type):
    """Apply formatting based on document type"""
    formatted = {}
    
    for section, sentences in content.items():
        if not sentences:
            continue
        
        section_text = ' '.join(sentences)
        
        formatted[section] = {
            'title': section.capitalize(),
            'content': section_text,
            'wordCount': len(word_tokenize(section_text)),
            'formatting': {
                'type': doc_type,
                'font': 'Arial',
                'fontSize': 11,
                'lineSpacing': 1.5,
                'margins': {'top': 1, 'right': 1, 'bottom': 1, 'left': 1}
            }
        }
    
    return formatted

def generate_document(structured_content, doc_type):
    """Generate document (DOCX format)"""
    try:
        doc = Document()
        
        # Set document properties
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add title
        doc.add_heading('DocStruct Generated Document', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Document Type: {doc_type}')
        doc.add_paragraph()
        
        # Add content sections
        for section_name, section_data in structured_content['formattedContent'].items():
            if section_data['content'] and section_data['content'].strip():
                # Add section heading
                doc.add_heading(section_data['title'], level=1)
                
                # Add content paragraphs
                content = section_data['content']
                paragraphs = split_into_paragraphs(content)
                
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                doc.add_paragraph()  # Add spacing between sections
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return {
            'buffer': buffer,
            'filename': f'DocStruct-{datetime.now().strftime("%Y%m%d%H%M%S")}.docx',
            'type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
    except Exception as e:
        logger.error(f"Document generation error: {str(e)}")
        raise

def split_into_paragraphs(text, max_words=50):
    """Split text into paragraphs based on word count"""
    words = text.split()
    paragraphs = []
    current_paragraph = []
    word_count = 0
    
    for word in words:
        current_paragraph.append(word)
        word_count += 1
        
        if word_count >= max_words:
            paragraphs.append(' '.join(current_paragraph))
            current_paragraph = []
            word_count = 0
    
    if current_paragraph:
        paragraphs.append(' '.join(current_paragraph))
    
    return paragraphs

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Route not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {str(error)}")
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3001, debug=False)